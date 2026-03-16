"""Unit tests for file parsers — PDF, DOCX, TXT parsing to ContentBlock."""

from __future__ import annotations

import io

import pytest

from src.ingestion.file_parser import SUPPORTED_EXTENSIONS, parse_file


class TestParseText:
    def test_plain_text(self) -> None:
        content = b"Hello, this is a test document.\n\nSecond paragraph."
        blocks = parse_file("test.txt", content)
        assert len(blocks) == 1
        assert blocks[0].is_text
        assert "Hello" in blocks[0].text

    def test_markdown_file(self) -> None:
        content = b"# Heading\n\nSome **bold** text."
        blocks = parse_file("readme.md", content)
        assert len(blocks) == 1
        assert "Heading" in blocks[0].text

    def test_empty_text_returns_empty(self) -> None:
        blocks = parse_file("empty.txt", b"")
        assert blocks == []

    def test_whitespace_only_returns_empty(self) -> None:
        blocks = parse_file("spaces.txt", b"   \n\n  ")
        assert blocks == []


class TestParseDocx:
    def test_simple_docx(self) -> None:
        from docx import Document

        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        buf = io.BytesIO()
        doc.save(buf)

        blocks = parse_file("test.docx", buf.getvalue())
        assert len(blocks) == 1
        assert "First paragraph" in blocks[0].text
        assert "Second paragraph" in blocks[0].text

    def test_empty_docx(self) -> None:
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)

        blocks = parse_file("empty.docx", buf.getvalue())
        assert blocks == []


class TestParsePdf:
    def test_simple_pdf(self) -> None:
        import pymupdf

        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello from PDF page 1")
        content = doc.tobytes()
        doc.close()

        blocks = parse_file("test.pdf", content)
        assert len(blocks) >= 1
        assert any("Hello from PDF" in (b.text or "") for b in blocks)

    def test_empty_pdf(self) -> None:
        import pymupdf

        doc = pymupdf.open()
        doc.new_page()  # blank page
        content = doc.tobytes()
        doc.close()

        blocks = parse_file("empty.pdf", content)
        assert blocks == []


class TestUnsupportedFormat:
    def test_unsupported_extension_raises(self) -> None:
        with pytest.raises(ValueError, match="Unsupported file type"):
            parse_file("data.csv", b"a,b,c")


class TestSupportedExtensions:
    def test_all_formats_listed(self) -> None:
        assert ".pdf" in SUPPORTED_EXTENSIONS
        assert ".docx" in SUPPORTED_EXTENSIONS
        assert ".txt" in SUPPORTED_EXTENSIONS
        assert ".md" in SUPPORTED_EXTENSIONS
